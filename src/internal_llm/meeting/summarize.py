from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from openai import OpenAI

from internal_llm.meeting.templates import SYSTEM_PROMPT, USER_PROMPT_TEMPLATE
from internal_llm.utils.config import get_config_value, load_project_config


def load_text_from_md_or_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def load_text_from_docx(path: Path, *, include_tables: bool = True) -> str:
    """Word(.docx) -> plain text with optional table rows."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx が未インストールです。pip install python-docx を実行してください。") from e

    doc = Document(str(path))
    lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    if include_tables and doc.tables:
        lines.append("\n[Tables]")
        for ti, table in enumerate(doc.tables, start=1):
            lines.append(f"(table {ti})")
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join((cell.text or "").split())
                    cells.append(cell_text)
                line = "\t".join([c for c in cells if c])
                if line.strip():
                    lines.append(line)

    return "\n".join(lines).strip()


def load_meeting_text(input_path: str, *, include_tables: bool = True) -> str:
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"Input not found: {path}")

    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return load_text_from_md_or_txt(path)
    if ext == ".docx":
        return load_text_from_docx(path, include_tables=include_tables)

    raise ValueError(f"Unsupported file type: {ext} (supported: .md, .txt, .docx)")


def call_llm_structuring(text: str, *, model: str) -> dict[str, Any]:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set (.env を確認してください)")

    client = OpenAI(api_key=api_key)
    user_prompt = USER_PROMPT_TEMPLATE.format(meeting_text=text)

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
    )

    content = resp.choices[0].message.content
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        cleaned = content.strip()
        cleaned = cleaned.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        return json.loads(cleaned)


def default_out_path(project_root: Path, input_path: str) -> Path:
    in_path = Path(input_path)
    paths_cfg = load_project_config(project_root, "configs/paths.yaml")
    outputs_dir = get_config_value(paths_cfg, ["paths", "outputs_dir"], "outputs")
    out_dir = project_root / str(outputs_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir / f"{in_path.stem}.summary.json"


def run_meeting_summary(
    *,
    project_root: Path,
    input_path: str,
    model: str,
    out_path: Path | None,
    include_tables: bool,
) -> Path:
    text = load_meeting_text(input_path, include_tables=include_tables)
    result = call_llm_structuring(text, model=model)

    target_path = out_path if out_path else default_out_path(project_root, input_path)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return target_path
