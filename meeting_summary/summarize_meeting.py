# meeting_summary/summarize_meeting.py
from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from openai import OpenAI

# --- ここが重要：直叩き実行でも import できるようにする ---
PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from meeting_summary.templates import SYSTEM_PROMPT, USER_PROMPT_TEMPLATE  # noqa: E402


def load_text_from_md_or_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def load_text_from_docx(path: Path, *, include_tables: bool = True) -> str:
    """
    Word(.docx) -> plain text
    - 段落を改行で連結
    - 表がある場合は簡易的に TSV 風に連結（任意）
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx が未インストールです。pip install python-docx を実行してください。") from e

    doc = Document(str(path))
    lines = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # tables (optional)
    if include_tables and doc.tables:
        lines.append("\n[Tables]")
        for ti, table in enumerate(doc.tables, start=1):
            lines.append(f"(table {ti})")
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join((cell.text or "").split())
                    cells.append(cell_text)
                # TSV 風
                line = "\t".join([c for c in cells if c])
                if line.strip():
                    lines.append(line)

    return "\n".join(lines).strip()


def load_meeting_text(input_path: str, *, include_tables: bool = True) -> str:
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"Input not found: {path}")

    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return load_text_from_md_or_txt(path)
    if ext == ".docx":
        return load_text_from_docx(path, include_tables=include_tables)

    raise ValueError(f"Unsupported file type: {ext} (supported: .md, .txt, .docx)")


def call_llm_structuring(text: str, *, model: str) -> dict:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set (.env を確認してください)")

    client = OpenAI(api_key=api_key)

    user_prompt = USER_PROMPT_TEMPLATE.format(meeting_text=text)

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
    )

    content = resp.choices[0].message.content

    # LLMがJSON文字列を返す想定
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        # たまに ```json ``` で囲まれる対策
        cleaned = content.strip()
        cleaned = cleaned.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        return json.loads(cleaned)


def default_out_path(input_path: str) -> Path:
    in_path = Path(input_path)
    out_dir = PROJECT_ROOT / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir / f"{in_path.stem}.summary.json"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="meetings/*.md or *.docx")
    parser.add_argument("--model", default="gpt-4.1-mini", help="chat model name")
    parser.add_argument("--out", default=None, help="output json path (optional)")
    parser.add_argument("--include_tables", action="store_true", help="include tables in docx extraction")
    args = parser.parse_args()

    text = load_meeting_text(args.input, include_tables=args.include_tables)

    result = call_llm_structuring(text, model=args.model)

    out_path = Path(args.out) if args.out else default_out_path(args.input)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
